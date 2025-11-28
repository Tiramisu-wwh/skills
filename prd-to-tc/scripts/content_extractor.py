#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
PRD文档内容提取器
支持Word和PDF文档的内容提取
"""

import sys
import os
import argparse
import json
from datetime import datetime

def extract_word_content(file_path):
    """提取Word文档内容"""
    try:
        from docx import Document
    except ImportError:
        print("错误: 需要安装python-docx库")
        print("运行: pip install python-docx")
        sys.exit(1)

    doc = Document(file_path)

    content = {
        'file_path': file_path,
        'extraction_time': datetime.now().isoformat(),
        'title': '',
        'sections': [],
        'tables': [],
        'full_text': ''
    }

    full_text = ""
    current_section = ""

    # 提取段落和标题
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        if para.style.name.startswith('Heading'):
            # 保存上一个section
            if current_section.strip():
                content['sections'].append(current_section.strip())

            level = para.style.name.split()[-1] if para.style.name.split()[-1].isdigit() else '1'
            section_header = f"{'#' * int(level)} {text}"
            current_section = section_header
            full_text += f"\n{section_header}\n"

            # 如果是一级标题，作为文档标题
            if level == '1' and not content['title']:
                content['title'] = text
        else:
            current_section += f"\n{text}"
            full_text += f"{text}\n"

    # 保存最后一个section
    if current_section.strip():
        content['sections'].append(current_section.strip())

    # 提取表格
    for i, table in enumerate(doc.tables):
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)

        content['tables'].append({
            'index': i + 1,
            'data': table_data
        })

        # 添加到全文
        full_text += f"\n\n## 表格 {i+1}\n"
        for row_idx, row in enumerate(table_data):
            full_text += " | ".join(row) + "\n"
            if row_idx == 0:
                full_text += "--- | " * len(row) + "\n"

    content['full_text'] = full_text.strip()

    return content

def extract_pdf_content(file_path):
    """提取PDF文档内容"""
    try:
        import PyPDF2
    except ImportError:
        print("错误: 需要安装PyPDF2库")
        print("运行: pip install PyPDF2")
        sys.exit(1)

    content = {
        'file_path': file_path,
        'extraction_time': datetime.now().isoformat(),
        'title': '',
        'sections': [],
        'tables': [],
        'full_text': ''
    }

    full_text = ""

    with open(file_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        for page_num, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text.strip():
                full_text += f"\n--- 第{page_num+1}页 ---\n{page_text.strip()}\n"

    content['full_text'] = full_text.strip()

    # 简单的标题提取
    lines = full_text.split('\n')
    for line in lines[:20]:  # 检查前20行
        line = line.strip()
        if line and len(line) < 100 and any(keyword in line for keyword in ['需求', 'PRD', '产品需求', '需求文档']):
            content['title'] = line
            break

    return content

def main():
    parser = argparse.ArgumentParser(description='PRD文档内容提取器')
    parser.add_argument('--word', type=str, help='Word文档路径')
    parser.add_argument('--pdf', type=str, help='PDF文档路径')
    parser.add_argument('--output', type=str, help='输出JSON文件路径（可选）')

    args = parser.parse_args()

    if not args.word and not args.pdf:
        print("错误: 必须指定 --word 或 --pdf 参数")
        parser.print_help()
        sys.exit(1)

    if args.word and args.pdf:
        print("错误: 不能同时指定 --word 和 --pdf 参数")
        sys.exit(1)

    file_path = args.word or args.pdf

    if not os.path.exists(file_path):
        print(f"错误: 文件不存在 {file_path}")
        sys.exit(1)

    try:
        # 提取内容
        if args.word:
            content = extract_word_content(file_path)
            file_type = "Word"
        else:
            content = extract_pdf_content(file_path)
            file_type = "PDF"

        print(f"✓ {file_type}文档内容提取完成")
        print(f"  - 标题: {content['title'] or '未识别'}")
        print(f"  - 章节: {len(content['sections'])}个")
        print(f"  - 表格: {len(content['tables'])}个")
        print(f"  - 总字符: {len(content['full_text'])}")

        # 输出结果
        if args.output:
            with open(args.output, 'w', encoding='utf-8') as f:
                json.dump(content, f, ensure_ascii=False, indent=2)
            print(f"✓ 内容已保存到: {args.output}")
        else:
            # 输出到标准输出
            print("\n=== 提取的内容 ===")
            print(json.dumps(content, ensure_ascii=False, indent=2))

    except Exception as e:
        print(f"错误: 内容提取失败 - {e}")
        sys.exit(1)

if __name__ == "__main__":
    main()